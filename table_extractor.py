"""
PDF / DOCX 表格检测与结构识别
- PDF: PyMuPDF 渲染 + TableTransformer 检测 + PaddleOCR 识别
- DOCX: python-docx 原生表格提取
"""
import io
import os
import fitz  # PyMuPDF
import config_data as config

# ========== 模型懒加载 ==========

_table_detector = None
_table_structurer = None
_table_ocr = None

_DETECTION_DIR = os.path.join(config.local_models_dir, "table-transformer-detection")
_STRUCTURE_DIR = os.path.join(config.local_models_dir, "table-transformer-structure-recognition")


def _get_table_detector():
    global _table_detector
    if _table_detector is None:
        from transformers import AutoImageProcessor, AutoModelForObjectDetection
        _table_detector = {
            "processor": AutoImageProcessor.from_pretrained(
                _DETECTION_DIR, local_files_only=True,
            ),
            "model": AutoModelForObjectDetection.from_pretrained(
                _DETECTION_DIR, local_files_only=True,
            ),
        }
    return _table_detector


def _get_table_structurer():
    global _table_structurer
    if _table_structurer is None:
        from transformers import AutoImageProcessor, AutoModelForObjectDetection
        _table_structurer = {
            "processor": AutoImageProcessor.from_pretrained(
                _STRUCTURE_DIR, local_files_only=True,
            ),
            "model": AutoModelForObjectDetection.from_pretrained(
                _STRUCTURE_DIR, local_files_only=True,
            ),
        }
    return _table_structurer


def _get_table_ocr():
    global _table_ocr
    if _table_ocr is None:
        from paddleocr import PaddleOCR
        _table_ocr = PaddleOCR(lang='ch')
    return _table_ocr


def preload_models():
    """预加载所有模型，避免首次请求时超时"""
    import sys
    sys.stderr.write("[Preload] 加载 TableTransformer + PaddleOCR...\n")
    sys.stderr.flush()
    _get_table_detector()
    sys.stderr.write("[Preload] TableTransformer detection 就绪\n")
    sys.stderr.flush()
    _get_table_structurer()
    sys.stderr.write("[Preload] TableTransformer structure 就绪\n")
    sys.stderr.flush()
    _get_table_ocr()
    sys.stderr.write("[Preload] PaddleOCR 就绪\n")
    sys.stderr.flush()


# ========== PDF 表格提取 ==========

def extract_tables_from_pdf(file_bytes: bytes, dpi: int = 150) -> list[dict]:
    """从 PDF 中提取所有表格，返回结构化数据

    Returns:
        [{"page_num": 1, "markdown": "| col1 | col2 |\n| --- | --- |\n| ..."}, ...]
    """
    import torch
    detector = _get_table_detector()
    structurer = _get_table_structurer()
    ocr = _get_table_ocr()

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    all_tables = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=dpi)
        img = pix.tobytes("png")

        from PIL import Image
        pil_img = Image.open(io.BytesIO(img)).convert("RGB")

        det_inputs = detector["processor"](images=pil_img, return_tensors="pt")
        with torch.no_grad():
            det_outputs = detector["model"](**det_inputs)

        tables = _parse_detections(det_outputs, pil_img.size, threshold=0.8)

        for table_idx, table_bbox in enumerate(tables):
            x1, y1, x2, y2 = [int(v) for v in table_bbox]
            cropped = pil_img.crop((x1, y1, x2, y2))

            str_inputs = structurer["processor"](images=cropped, return_tensors="pt")
            with torch.no_grad():
                str_outputs = structurer["model"](**str_inputs)

            cells = _parse_structure(str_outputs, cropped.size, threshold=0.7)

            if not cells:
                continue

            rows = _cells_to_rows(cells, cropped)
            markdown = _rows_to_markdown(rows, ocr, cropped)

            all_tables.append({
                "page_num": page_num + 1,
                "table_index": table_idx,
                "markdown": markdown,
            })

    doc.close()
    return all_tables


def _parse_detections(outputs, img_size, threshold=0.8) -> list:
    import torch
    w, h = img_size
    results = []
    for score, label, box in zip(
        outputs["scores"], outputs["labels"], outputs["boxes"]
    ):
        if score < threshold:
            continue
        box = [round(c.item(), 2) for c in box]
        xc, yc, bw, bh = box
        x1 = (xc - bw / 2) * w
        y1 = (yc - bh / 2) * h
        x2 = (xc + bw / 2) * w
        y2 = (yc + bh / 2) * h
        results.append((x1, y1, x2, y2))
    return results


def _parse_structure(outputs, img_size, threshold=0.7) -> list[dict]:
    w, h = img_size
    cells = []
    for score, label_id, box in zip(
        outputs["scores"], outputs["labels"], outputs["boxes"]
    ):
        if score < threshold:
            continue
        box = [round(c.item(), 2) for c in box]
        xc, yc, bw, bh = box
        x1 = max(0, (xc - bw / 2) * w)
        y1 = max(0, (yc - bh / 2) * h)
        x2 = min(w, (xc + bw / 2) * w)
        y2 = min(h, (yc + bh / 2) * h)
        cells.append({
            "bbox": (x1, y1, x2, y2),
            "label": label_id.item(),
        })
    return cells


def _cells_to_rows(cells: list[dict], cropped_img) -> list[list[tuple]]:
    if not cells:
        return []

    cells_sorted = sorted(cells, key=lambda c: c["bbox"][1])
    rows = []
    current_row = [cells_sorted[0]]
    row_y_max = cells_sorted[0]["bbox"][3]

    for cell in cells_sorted[1:]:
        y1 = cell["bbox"][1]
        if y1 > row_y_max + 5:
            rows.append(sorted(current_row, key=lambda c: c["bbox"][0]))
            current_row = [cell]
            row_y_max = cell["bbox"][3]
        else:
            current_row.append(cell)
            row_y_max = max(row_y_max, cell["bbox"][3])

    if current_row:
        rows.append(sorted(current_row, key=lambda c: c["bbox"][0]))

    return rows


def _rows_to_markdown(rows, ocr, cropped_img) -> str:
    from PIL import Image

    markdown_rows = []
    for row_cells in rows:
        texts = []
        for cell in row_cells:
            x1, y1, x2, y2 = [int(v) for v in cell["bbox"]]
            cell_img = cropped_img.crop((x1, y1, x2, y2))
            if cell_img.width < 5 or cell_img.height < 5:
                texts.append("")
                continue
            import numpy as np
            result = list(ocr.ocr(np.array(cell_img)))
            if result and result[0]:
                rec_texts = result[0].get("rec_texts", [])
                cell_text = " ".join(rec_texts)
                texts.append(cell_text.replace("|", "/"))
            else:
                texts.append("")
        markdown_rows.append(texts)

    if not markdown_rows:
        return ""

    max_cols = max(len(row) for row in markdown_rows)
    lines = []
    for i, row in enumerate(markdown_rows):
        while len(row) < max_cols:
            row.append("")
        lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    return "\n".join(lines)


# ========== DOCX 表格提取 ==========

def extract_tables_from_docx(file_bytes: bytes) -> list[str]:
    """从 DOCX 中提取所有表格，返回 Markdown 格式列表"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.replace("|", "/").replace("\n", " ") for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(_rows_to_markdown_str(rows))
    return tables


def _rows_to_markdown_str(rows: list[list[str]]) -> str:
    """将行列数据转为 Markdown 表格字符串"""
    max_cols = max(len(row) for row in rows)
    lines = []
    for i, row in enumerate(rows):
        while len(row) < max_cols:
            row.append("")
        lines.append("| " + " | ".join(row) + " |")
        if i == 0:
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    return "\n".join(lines)


def _get_heading_level(style_name: str) -> int:
    """从 Word 样式名提取标题层级，非标题返回 0"""
    import re
    if not style_name:
        return 0
    m = re.match(r'(?:heading|标题)\s*(\d)', style_name.lower().strip())
    if m:
        level = int(m.group(1))
        return level if 1 <= level <= 3 else 0
    return 0


def read_docx_with_tables(file_bytes: bytes) -> str:
    """读取 DOCX 全文，标题转为 markdown 格式，表格原位转为 Markdown 表格"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(file_bytes))
    tables_md = extract_tables_from_docx(file_bytes)

    parts = []
    para_idx = 0
    table_md_idx = 0

    body = doc.element.body
    for child in body:
        if child.tag == qn('w:p'):
            if para_idx >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_idx]
            para_idx += 1
            text = para.text.strip()
            if not text:
                continue
            level = _get_heading_level(para.style.name if para.style else '')
            if level > 0:
                parts.append('#' * level + ' ' + text)
            else:
                parts.append(text)
        elif child.tag == qn('w:tbl'):
            if table_md_idx < len(tables_md):
                parts.append('\n' + tables_md[table_md_idx] + '\n')
                table_md_idx += 1

    return '\n'.join(parts)
